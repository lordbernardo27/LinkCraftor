from __future__ import annotations

import asyncio
import io
import json
import shutil
import tempfile

from pathlib import Path

from docx import Document
from fastapi import UploadFile

import backend.server.routes.files as files_route

from backend.server.pipelines.upload_document.uploaded_document_to_uduc_pipeline.upload_intake import (
    UploadIntakeDependencies,
    run_upload_intake,
)


results = []


def check(name: str, condition: bool) -> None:
    status = "PASS" if condition else "FAIL"
    results.append((name, status))
    print(f"{name}: {status}")


def make_upload(filename: str, raw: bytes) -> UploadFile:
    return UploadFile(
        filename=filename,
        file=io.BytesIO(raw),
    )


def make_docx_bytes(marker: str) -> bytes:
    buffer = io.BytesIO()

    document = Document()
    document.add_heading("U3.13 DOCX Heading", level=1)
    document.add_paragraph(marker)

    document.save(buffer)
    return buffer.getvalue()


def normalize_index_rows(payload):
    if isinstance(payload, list):
        return payload

    if isinstance(payload, dict):
        items = payload.get("items")

        if isinstance(items, list):
            return items

    return []


async def main() -> None:
    temp_root = Path(
        tempfile.mkdtemp(
            prefix="linkcraftor_u3_13_success_"
        )
    )

    docs_root = temp_root / "docs"

    original_data_dir = files_route.DATA_DIR
    original_docs_dir = files_route.DOCS_DIR

    try:
        # ----------------------------------------------------
        # Strict isolation from production upload data.
        # ----------------------------------------------------

        files_route.DATA_DIR = temp_root
        files_route.DOCS_DIR = docs_root

        deps = UploadIntakeDependencies(
            guess_extension=files_route._guess_ext,
            normalize_workspace_id=files_route._ws,
            extract_preview=files_route._extract_preview_from_bytes,
            store_and_index=files_route._store_and_index,
            rollback_committed_upload=files_route._rollback_committed_upload,
            workspace_directory=files_route._ws_dir,
            allowed_extensions=files_route.ALLOWED_EXT,
        )

        marker_txt = "U3_13_CANONICAL_TXT_MARKER"
        marker_md = "U3_13_CANONICAL_MD_MARKER"
        marker_markdown = "U3_13_CANONICAL_MARKDOWN_MARKER"
        marker_html = "U3_13_CANONICAL_HTML_MARKER"
        marker_htm = "U3_13_CANONICAL_HTM_MARKER"
        marker_docx = "U3_13_CANONICAL_DOCX_MARKER"

        cases = [
            (
                "TXT",
                "canonical.txt",
                (
                    "U3.13 TXT Heading\n\n"
                    + marker_txt
                    + "\nCanonical TXT body."
                ).encode("utf-8"),
                marker_txt,
            ),
            (
                "MD",
                "canonical.md",
                (
                    "# U3.13 Markdown Heading\n\n"
                    + marker_md
                    + "\n\nCanonical Markdown body."
                ).encode("utf-8"),
                marker_md,
            ),
            (
                "MARKDOWN",
                "canonical.markdown",
                (
                    "# U3.13 Markdown Alias Heading\n\n"
                    + marker_markdown
                    + "\n\nCanonical .markdown body."
                ).encode("utf-8"),
                marker_markdown,
            ),
            (
                "HTML",
                "canonical.html",
                (
                    "<html><body>"
                    "<h1>U3.13 HTML Heading</h1>"
                    "<p>"
                    + marker_html
                    + "</p>"
                    "<p>Canonical HTML body.</p>"
                    "</body></html>"
                ).encode("utf-8"),
                marker_html,
            ),
            (
                "HTM",
                "canonical.htm",
                (
                    "<html><body>"
                    "<h1>U3.13 HTM Heading</h1>"
                    "<p>"
                    + marker_htm
                    + "</p>"
                    "<p>Canonical HTM body.</p>"
                    "</body></html>"
                ).encode("utf-8"),
                marker_htm,
            ),
            (
                "DOCX",
                "canonical.docx",
                make_docx_bytes(marker_docx),
                marker_docx,
            ),
        ]

        required_extraction_fields = {
            "source_path",
            "source_type",
            "title",
            "text",
            "headings",
            "metadata",
            "extraction_status",
            "extraction_confidence",
            "created_at",
        }

        created_document_ids = []

        for label, filename, raw, marker in cases:
            print()
            print("----------------------------------------")
            print(f"FORMAT: {label}")
            print("----------------------------------------")

            workspace_id = f"ws_u3_13_{label.lower()}"

            result = await run_upload_intake(
                workspace_id=workspace_id,
                file=make_upload(filename, raw),
                dependencies=deps,
            )

            check(
                f"{label}_INTAKE_OK",
                result.get("ok") is True,
            )

            check(
                f"{label}_WORKSPACE_PRESERVED",
                result.get("workspace_id") == workspace_id,
            )

            doc = result.get("doc") or {}
            extraction = result.get("extraction") or {}

            document_id = str(
                doc.get("doc_id")
                or doc.get("document_id")
                or ""
            ).strip()

            stored_name = str(
                doc.get("stored_name")
                or ""
            ).strip()

            created_document_ids.append(document_id)

            check(
                f"{label}_DOCUMENT_ID_CREATED",
                bool(document_id),
            )

            check(
                f"{label}_ORIGINAL_FILENAME_PRESERVED",
                doc.get("filename") == filename,
            )

            check(
                f"{label}_STORED_FILENAME_CREATED",
                bool(stored_name)
                and stored_name.startswith(
                    document_id + "__"
                ),
            )

            persisted_source = (
                files_route._ws_dir(workspace_id)
                / stored_name
            )

            check(
                f"{label}_SOURCE_FILE_PERSISTED",
                persisted_source.is_file(),
            )

            check(
                f"{label}_SOURCE_BYTES_EXACT",
                persisted_source.is_file()
                and persisted_source.read_bytes() == raw,
            )

            check(
                f"{label}_EXTRACTION_RESULT_FIELDS_PRESENT",
                isinstance(extraction, dict)
                and required_extraction_fields.issubset(
                    extraction.keys()
                ),
            )

            check(
                f"{label}_EXTRACTION_STATUS_SUCCESS",
                extraction.get("extraction_status")
                == "success",
            )

            check(
                f"{label}_SOURCE_TYPE_PRESENT",
                bool(
                    str(
                        extraction.get("source_type")
                        or ""
                    ).strip()
                ),
            )

            extraction_text = str(
                extraction.get("text")
                or ""
            )

            check(
                f"{label}_CANONICAL_MARKER_EXTRACTED",
                marker in extraction_text,
            )

            extraction_source_path = str(
                extraction.get("source_path")
                or ""
            ).strip()

            expected_source_path = str(
                persisted_source.resolve()
            )

            check(
                f"{label}_EXTRACTION_SOURCE_PATH_IS_PERSISTED_SOURCE",
                bool(extraction_source_path)
                and str(
                    Path(extraction_source_path).resolve()
                )
                == expected_source_path,
            )

            index_path = (
                files_route._ws_dir(workspace_id)
                / "index.json"
            )

            check(
                f"{label}_REGISTRY_EXISTS",
                index_path.is_file(),
            )

            rows = []

            if index_path.is_file():
                raw_index = json.loads(
                    index_path.read_text(
                        encoding="utf-8"
                    )
                )

                rows = normalize_index_rows(
                    raw_index
                )

            matching_rows = [
                row
                for row in rows
                if isinstance(row, dict)
                and row.get("doc_id") == document_id
            ]

            check(
                f"{label}_REGISTRY_HAS_EXACT_DOCUMENT",
                len(matching_rows) == 1,
            )

            check(
                f"{label}_REGISTRY_STORED_NAME_MATCHES",
                len(matching_rows) == 1
                and matching_rows[0].get(
                    "stored_name"
                )
                == stored_name,
            )

            check(
                f"{label}_JOB_ID_NONE",
                result.get("job_id") is None,
            )

            check(
                f"{label}_PROCESSING_STATUS_NOT_APPLICABLE",
                result.get("processing_status")
                == "not_applicable",
            )

        print()
        print("----------------------------------------")
        print("CROSS-FORMAT IDENTITY")
        print("----------------------------------------")

        check(
            "ALL_SIX_DOCUMENT_IDS_CREATED",
            len(created_document_ids) == 6
            and all(created_document_ids),
        )

        check(
            "ALL_SIX_DOCUMENT_IDS_UNIQUE",
            len(set(created_document_ids)) == 6,
        )

    finally:
        files_route.DATA_DIR = original_data_dir
        files_route.DOCS_DIR = original_docs_dir

        temp_prefix = str(
            temp_root.resolve()
        )

        with files_route._INDEX_LOCKS_GUARD:
            stale_keys = [
                key
                for key in files_route._INDEX_LOCKS
                if key.startswith(temp_prefix)
            ]

            for key in stale_keys:
                files_route._INDEX_LOCKS.pop(
                    key,
                    None,
                )

        shutil.rmtree(
            temp_root,
            ignore_errors=True,
        )

    check(
        "TEMP_TEST_ROOT_REMOVED",
        not temp_root.exists(),
    )

    failures = [
        name
        for name, status in results
        if status != "PASS"
    ]

    print()
    print("========================================")

    if failures:
        print(
            "U3.13_INTAKE_SUCCESS_FORMATS_VERIFICATION: FAIL"
        )

        print("FAILED_CHECKS:")

        for failure in failures:
            print(f" - {failure}")

        raise RuntimeError(
            "U3.13 success-format verification failed."
        )

    print(
        "U3.13_INTAKE_SUCCESS_FORMATS_VERIFICATION: PASS"
    )


if __name__ == "__main__":
    asyncio.run(main())
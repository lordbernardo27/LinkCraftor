from __future__ import annotations
# backend/server/main.py


# -------------------------------


# FastAPI backend exposing:


#  - GET  /health + /api/health  (via health_router)


#  - POST /api/convert/docx


#  - POST /api/export/docx


#  - GET  /api/external/resolve


#  - POST /api/external/log


#  - FRONTEND served at /


#  - Static assets at /assets/*


#  - Static UI pages at /static/*


#  - Business Dashboard at /business/*


#


# Uses python-mammoth to convert .docx -> HTML.







import asyncio


from contextlib import asynccontextmanager





import io


import logging


from pathlib import Path





import mammoth


from backend.server.owner.audit_routes import router as owner_audit_router


from fastapi import Body, FastAPI, File, HTTPException, UploadFile


from fastapi.middleware.cors import CORSMiddleware


from fastapi.responses import FileResponse, RedirectResponse, Response, StreamingResponse


from fastapi.staticfiles import StaticFiles





# =========================


# Router imports


# =========================


from .routes.imported_targets_urls_compat import router as imported_urls_router


from .routes.draft_topics import router as draft_router


from .routes.planning import router as planning_router


from .routes.engine_scoring import router as engine_scoring_router


from .routes.files import router as files_router


from .routes.health import router as health_router


from .routes.engine_run import router as engine_run_router


from .routes.engine_decisions import router as engine_decisions_router





from backend.server.routes.rb2_run import router as rb2_runner_router


from backend.server.routes.document_registry import router as document_registry_router


from backend.server.routes.workspace_health import router as workspace_health_router


from backend.server.routes.rebuild_routes import router as rebuild_router


from backend.server.routes.reload_routes import router as reload_router


from backend.server.orchestration.routes import router as orchestration_router


from backend.server.tms.routes import router as tms_router


from backend.server.routes.workspace_autosave import router as workspace_autosave_router
from backend.server.routes.external.runtime import router as external_runtime_router
from backend.server.routes.external.manual import router as external_manual_router
from backend.server.routes.external.auto import router as external_auto_router
from backend.server.routes.external.owner_sources import router as external_owner_sources_router
from backend.server.routes.external.resolver import router as resolver_runtime_router
from backend.server.routes.external.manual import router as manual_runtime_router
from backend.server.routes.external.sources import router as sources_runtime_router
from backend.server.routes.external.import_runs import router as import_runs_runtime_router
from backend.server.routes.external.auto_cleanup import router as auto_cleanup_runtime_router





log = logging.getLogger("linkcraftor.server")


logging.basicConfig(level=logging.INFO)





app = FastAPI(title="LinkCraftor API", version="0.1.0")





app.include_router(owner_audit_router)


app.include_router(workspace_autosave_router)








# =========================


# Path helpers


# =========================


def _normalize_prefix(p: str) -> str:


    p = (p or "").strip()


    if not p:


        return ""


    if not p.startswith("/"):


        p = "/" + p


    return p.rstrip("/") or "/"








def _already_mounted(module_path: str, prefix: str) -> bool:


    pfx = _normalize_prefix(prefix)


    for r in app.routes:


        path = getattr(r, "path", None) or getattr(r, "path_format", "")


        endpoint = getattr(r, "endpoint", None)


        mod = getattr(endpoint, "__module__", "") if endpoint else ""


        if mod.startswith(module_path) and str(path).startswith(pfx):


            return True


    return False








BASE_DIR = Path(__file__).resolve().parent


STATIC_DIR = BASE_DIR / "static"


BUSINESS_DIR = STATIC_DIR / "business"








# =========================


# Static mounts


# =========================


if STATIC_DIR.is_dir():


    app.mount("/static", StaticFiles(directory=str(STATIC_DIR), html=True), name="static")





if BUSINESS_DIR.is_dir():


    app.mount("/business", StaticFiles(directory=str(BUSINESS_DIR), html=True), name="business")





OWNER_DIR = Path("backend/server/owner")


if OWNER_DIR.is_dir():


    app.mount("/owner", StaticFiles(directory=str(OWNER_DIR), html=True), name="owner")








@app.get("/business", include_in_schema=False)


def business_root():


    return RedirectResponse(url="/business/")








# =========================


# CORS


# =========================


app.add_middleware(


    CORSMiddleware,


    allow_origins=["*"],


    allow_credentials=True,


    allow_methods=["*"],


    allow_headers=["*"],


)








# =========================


# Frontend serving


# =========================


FRONTEND_DIR = (BASE_DIR / ".." / ".." / "frontend" / "public").resolve()


ASSETS_DIR = FRONTEND_DIR / "assets"





if ASSETS_DIR.is_dir():


    app.mount("/assets", StaticFiles(directory=str(ASSETS_DIR), html=False), name="assets")








@app.get("/")


def serve_index():


    index_path = FRONTEND_DIR / "index.html"


    if not index_path.is_file():


        raise HTTPException(status_code=404, detail=f"Missing frontend file: {index_path}")


    return FileResponse(str(index_path))








@app.get("/favicon.ico", include_in_schema=False)


def favicon():


    icon_path = STATIC_DIR / "favicon.ico"


    if icon_path.is_file():


        return FileResponse(str(icon_path))


    return Response(status_code=204)








# =========================


# DOCX Ã¢â€ â€™ HTML converter


# =========================


@app.post("/api/convert/docx")


async def convert_docx(file: UploadFile = File(...)):


    if file is None or not (file.filename or "").lower().endswith(".docx"):


        raise HTTPException(status_code=400, detail="Please upload a .docx file.")





    try:


        data = await file.read()





        with io.BytesIO(data) as buff:


            result = mammoth.convert_to_html(buff)


            html = result.value or ""





        try:


            with io.BytesIO(data) as buff2:


                raw = mammoth.extract_raw_text(buff2)


                text = raw.value or ""


        except Exception:


            text = ""





        return {"ok": True, "filename": file.filename, "ext": ".docx", "html": html, "text": text}


    except Exception as e:


        raise HTTPException(status_code=500, detail=f"DOCX conversion error: {e}")








# ================================


# Export simple HTML Ã¢â€ â€™ .docx


# ================================


@app.post("/api/export/docx")


async def export_docx(payload: dict = Body(...)):


    html = (payload or {}).get("html", "")


    filename = (payload or {}).get("filename", "export.docx")





    if not html or not str(html).strip():


        raise HTTPException(status_code=400, detail="html is required")





    try:


        from io import BytesIO


        from bs4 import BeautifulSoup


        from docx import Document


    except Exception as ie:


        raise HTTPException(


            status_code=500,


            detail=(


                f"Export dependencies missing: {ie}. "


                "Install with: pip install python-docx beautifulsoup4"


            ),


        )





    soup = BeautifulSoup(html, "html.parser")


    root = soup.body or soup


    doc = Document()





    def add_text_block(text: str):


        t = (text or "").strip()


        if t:


            doc.add_paragraph(t)





    for el in getattr(root, "children", []):


        if getattr(el, "name", None) is None:


            add_text_block(str(el))


            continue





        name = (el.name or "").lower()


        text = el.get_text(" ", strip=True)





        if name in ("h1", "h2", "h3", "h4", "h5", "h6"):


            level = min(int(name[1]), 9)


            doc.add_heading(text, level=level)


        elif name == "p":


            doc.add_paragraph(text)


        elif name in ("ul", "ol"):


            bullet = (name == "ul")


            for li in el.find_all("li", recursive=False):


                p = doc.add_paragraph(li.get_text(" ", strip=True))


                p.style = "List Bullet" if bullet else "List Number"


        elif name == "br":


            doc.add_paragraph("")


        else:


            if text:


                doc.add_paragraph(text)





    if not filename.lower().endswith(".docx"):


        filename += ".docx"





    buf = BytesIO()


    doc.save(buf)


    buf.seek(0)





    return StreamingResponse(


        buf,


        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",


        headers={"Content-Disposition": f'attachment; filename="{filename}"'},


    )








# =========================


# Optional Site Reader


# =========================


site_reader_mount_error = None


site_reader_router = None


try:


    from backend.server.routes.site_reader import router as site_reader_router


except Exception as e:


    site_reader_router = None


    site_reader_mount_error = repr(e)








# =========================


# Core API mounts


# =========================


if not _already_mounted("backend.server.routes.health", "/"):


    app.include_router(health_router, tags=["health"])





if not _already_mounted("backend.server.routes.imported_targets_urls_compat", "/api/urls"):


    app.include_router(imported_urls_router, tags=["urls"])





if not _already_mounted("backend.server.routes.draft_topics", "/api/draft"):


    app.include_router(draft_router, tags=["draft"])





if not _already_mounted("backend.server.routes.planning", "/api/planning"):


    app.include_router(planning_router, tags=["planning"])





if not _already_mounted("backend.server.routes.engine_run", "/api/engine"):


    app.include_router(engine_run_router, tags=["engine"])





if not _already_mounted("backend.server.routes.engine_decisions", "/api/engine"):


    app.include_router(engine_decisions_router, tags=["engine-decisions"])





if not _already_mounted("backend.server.routes.engine_scoring", "/api/engine"):


    app.include_router(engine_scoring_router, tags=["engine-scoring"])





if not _already_mounted("backend.server.routes.files", "/api"):


    app.include_router(files_router, tags=["files"])





if not _already_mounted("backend.server.routes.rb2_run", "/api/rb2"):


    app.include_router(rb2_runner_router, prefix="/api/rb2", tags=["rb2"])





if not _already_mounted("backend.server.routes.workspace_health", "/api/workspace"):


    app.include_router(workspace_health_router, tags=["workspace"])





if not _already_mounted("backend.server.routes.rebuild_routes", "/api/rebuild"):


    app.include_router(rebuild_router)





if not _already_mounted("backend.server.routes.reload_routes", "/api/reload"):


    app.include_router(reload_router)





app.include_router(orchestration_router)





if not _already_mounted("backend.server.routes.document_registry", "/api/site/target_pools/document_registry"):


    app.include_router(


        document_registry_router,


        prefix="/api/site/target_pools/document_registry",


        tags=["document-registry"],


    )





if not _already_mounted("backend.server.tms.routes", "/api/tms"):


    app.include_router(tms_router)





if not _already_mounted("backend.server.routes.external.owner_sources", "/api/external/owner_sources_runtime"):
    app.include_router(
        external_owner_sources_router,
        prefix="/api/external/owner_sources_runtime",
        tags=["external-owner-sources"],
    )
if not _already_mounted("backend.server.routes.external.auto", "/api/external/auto_runtime"):
    app.include_router(
        external_auto_router,
        prefix="/api/external/auto_runtime",
        tags=["external-auto"],
    )
if not _already_mounted("backend.server.routes.external.manual", "/api/external/manual_runtime"):
    app.include_router(
        external_manual_router,
        prefix="/api/external/manual_runtime",
        tags=["external-manual"],
    )
if not _already_mounted("backend.server.routes.external.runtime", "/api/external/runtime"):
    app.include_router(
        external_runtime_router,
        prefix="/api/external/runtime",
        tags=["external-runtime"],
    )
if site_reader_router is not None:


    if not _already_mounted("backend.server.routes.site_reader", "/api/site"):


        app.include_router(site_reader_router, prefix="/api/site", tags=["site-reader"])
app.include_router(resolver_runtime_router, prefix="/api/external/resolver_runtime")








# =========================


# Debug route list


# =========================


@app.get("/__routes")


def routes():


    out = []


    for r in app.router.routes:


        methods = sorted(list(getattr(r, "methods", []) or []))


        path = getattr(r, "path", None)


        if path:


            out.append({"path": path, "methods": methods})


    return {"routes": out, "site_reader_mount_error": site_reader_mount_error}






if not _already_mounted("backend.server.routes.external.sources", "/api/external/sources_runtime"):
    app.include_router(
        sources_runtime_router,
        prefix="/api/external/sources_runtime",
    )


if not _already_mounted("backend.server.routes.external.import_runs", "/api/external/import_runs_runtime"):
    app.include_router(
        import_runs_runtime_router,
        prefix="/api/external/import_runs_runtime",
    )


if not _already_mounted("backend.server.routes.external.auto_cleanup", "/api/external/auto_cleanup_runtime"):
    app.include_router(
        auto_cleanup_runtime_router,
        prefix="/api/external/auto_cleanup_runtime",
    )


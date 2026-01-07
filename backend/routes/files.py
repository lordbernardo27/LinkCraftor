# backend/routes/files.py

import os
import io
import re
from html import unescape
from flask import Blueprint, request, jsonify, send_file
from backend.services.storage_local import LocalStorage

# --- Blueprint MUST be defined before any @bp.route decorators ---
bp = Blueprint("files", __name__)

# Simple storage singleton
_storage: "LocalStorage | None" = None
def storage() -> LocalStorage:
    global _storage
    if _storage is None:
        root = os.path.join(os.getcwd(), "data")
        _storage = LocalStorage(root)
    return _storage

# ---------- helpers ----------
def _ext(fn: str) -> str:
    return os.path.splitext(fn)[1].lower()

def _decode_utf8(b: bytes) -> str:
    try:
        return b.decode("utf-8", errors="ignore")
    except Exception:
        return ""

def _docx_to_html(data: bytes) -> str:
    """
    Very light DOCX -> HTML for preview (paragraphs only).
    Links created in the editor/export pipeline later.
    """
    from docx import Document  # python-docx
    f = io.BytesIO(data)
    doc = Document(f)
    out = []
    for p in doc.paragraphs:
        t = p.text or ""
        t = unescape(t)
        # escape minimal HTML special chars
        t = (t.replace("&", "&amp;")
               .replace("<", "&lt;")
               .replace(">", "&gt;"))
        out.append(f"<p>{t}</p>")
    return "\n".join(out) if out else "<p></p>"

# ---------- routes ----------

@bp.post("/upload")
def upload():
    """
    Accepts a single file field named 'file'.
    Returns JSON: { ok, filename, ext, text? or html? } and saves the original.
    """
    if "file" not in request.files:
        return jsonify({"ok": False, "error": "No file field 'file'"}), 400

    f = request.files["file"]
    if not f.filename:
        return jsonify({"ok": False, "error": "Empty filename"}), 400

    filename = f.filename
    ext = _ext(filename)
    data = f.read()

    # Save original
    saved_name = storage().save_original(filename, data)

    # Supported preview types
    if ext in (".html", ".htm"):
        payload = {"ok": True, "filename": saved_name, "ext": ext, "html": _decode_utf8(data)}
        return jsonify(payload)

    if ext in (".txt", ".md"):
        payload = {"ok": True, "filename": saved_name, "ext": ext, "text": _decode_utf8(data)}
        return jsonify(payload)

    if ext == ".docx":
        try:
            html_preview = _docx_to_html(data)
        except Exception as e:
            return jsonify({"ok": False, "error": f"DOCX parse failed: {e}"}), 400
        payload = {"ok": True, "filename": saved_name, "ext": ext, "html": html_preview}
        return jsonify(payload)

    return jsonify({
        "ok": False,
        "error": f"Unsupported file type '{ext}'. Use .docx, .md, .html, .htm, .txt.",
        "filename": saved_name,
        "ext": ext
    }), 400


@bp.get("/download")
def download_original():
    fname = request.args.get("filename") or ""
    if not fname:
        return jsonify({"ok": False, "error": "filename query param required"}), 400
    path = storage().get_original_path(fname)
    if not os.path.isfile(path):
        return jsonify({"ok": False, "error": "File not found"}), 404
    return send_file(path, as_attachment=True, download_name=fname)


@bp.route("/export/docx", methods=["POST"])
def export_docx():
    """
    Receives JSON: { filename, content } where content is editor HTML.
    Builds a .docx and preserves clickable links for:
      - <a href="...">text</a>
      - <span class="lc-underlined" data-url="...">text</span>  (internal tool’s underline)
    """
    from html.parser import HTMLParser
    from docx import Document
    from docx.shared import RGBColor
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement        # type: ignore[reportPrivateImportUsage]
    from docx.oxml.ns import qn

    data = request.get_json(silent=True) or {}
    base = (data.get("filename") or "document").rsplit(".", 1)[0]
    html_in = data.get("content") or ""
    out_name = f"{base}.docx"
    out_path = storage().export_path(out_name)

    def looks_http(u: str) -> bool:
        return bool(re.match(r"^https?://", str(u or ""), re.I))

    def add_hyperlink(paragraph, url, text):
        """
        Append a clickable hyperlink run to `paragraph` (blue & underlined).
        """
        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

        # <w:hyperlink r:id="..."><w:r><w:rPr>...</w:rPr><w:t>text</w:t></w:r></w:hyperlink>
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0000FF")
        rPr.append(color)

        new_run.append(rPr)

        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")  # keep spaces if any
        t.text = text
        new_run.append(t)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    class MiniHTMLToDocx(HTMLParser):
        """
        Minimal HTML → DOCX converter:
          - new paragraph for <p>, <div>, and headings <h1..h3>
          - preserves <br>
          - clickable for <a href> and <span class="lc-underlined" data-url>
          - plain text otherwise
        """
        def __init__(self, document):
            super().__init__(convert_charrefs=True)
            self.doc = document
            self.p = None
            self.current_link = None

        def _ensure_par(self, style=None):
            if self.p is None:
                self.p = self.doc.add_paragraph()
                if style:
                    try:
                        self.p.style = style
                    except Exception:
                        pass
            return self.p

        def handle_starttag(self, tag, attrs):
            a = dict(attrs)
            if tag in ("p", "div"):
                self.p = self.doc.add_paragraph()
            elif tag in ("h1", "h2", "h3"):
                self.p = self.doc.add_paragraph()
                try:
                    self.p.style = f"Heading {tag[1]}"
                except Exception:
                    pass
            elif tag == "br":
                self._ensure_par().add_run().add_break()
            elif tag == "a":
                self.current_link = a.get("href") or ""
            elif tag == "span":
                cls = (a.get("class") or "").split()
                if "lc-underlined" in cls:
                    self.current_link = a.get("data-url") or ""

        def handle_endtag(self, tag):
            if tag in ("p", "div", "h1", "h2", "h3"):
                self.p = None
            elif tag in ("a", "span"):
                self.current_link = None

        def handle_data(self, data):
            text = data or ""
            if not text:
                return
            par = self._ensure_par()
            if self.current_link and looks_http(self.current_link):
                add_hyperlink(par, self.current_link, text)
            elif self.current_link:
                # Non-http (e.g. "#section") → render as blue/underlined text
                run = par.add_run(text)
                run.font.underline = True
                run.font.color.rgb = RGBColor(0, 0, 255)
            else:
                par.add_run(text)

    # Build the document
    doc = Document()
    parser = MiniHTMLToDocx(doc)
    parser.feed(html_in)
    parser.close()
    doc.save(out_path)

    return send_file(
        out_path,
        as_attachment=True,
        download_name=out_name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@bp.get("/export/zip")
def export_zip():
    """Zip all originals in data/uploads and return."""
    import zipfile
    mem = io.BytesIO()
    with zipfile.ZipFile(mem, "w", zipfile.ZIP_DEFLATED) as z:
        for p, name in storage().iter_originals():
            z.write(p, arcname=name)
    mem.seek(0)
    return send_file(mem, as_attachment=True, download_name="linkcraftor_originals.zip", mimetype="application/zip")


@bp.get("/export/rar")
def export_rar():
    return jsonify({"ok": False, "error": "RAR export not supported. Use ZIP."}), 501
